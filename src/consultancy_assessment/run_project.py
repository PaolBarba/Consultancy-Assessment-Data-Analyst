"""
Generate a DOCX report summarizing the consultancy assessment.

Includes methodology, assumptions, results, and a plot image.
"""

from docx import Document
from docx.shared import Inches

from consultancy_assessment.scripts.assessment import ConsultancyAssessment


def generate_docx_report(image_path, output_path="report.docx"):
    """
    Generate a DOCX report summarizing the consultancy assessment, including methodology, assumptions, results table, and a plot image.

    Parameters
    ----------
    image_path : str
        Path to the image file to be included in the report.
    output_path : str, optional
        Path where the generated DOCX report will be saved (default is "report.docx").
    """
    doc = Document()
    assessment = ConsultancyAssessment()
    table_data = assessment.run()

    doc.add_heading("Consultancy Assessment Report", 0)

    doc.add_heading("Methodology and Assumptions", level=1)

    doc.add_paragraph(
        "The consultancy assessment was conducted through a structured data pipeline designed to process, clean, "
        "merge, analyze, and visualize global health indicators, specifically related to Under-5 Mortality Rate (U5MR) "
        "status and population data. Below is a breakdown of the methodology followed, along with key assumptions and caveats:"
    )

    doc.add_paragraph("- **Data Loading:** Three key datasets were used:\n"
                    "  - U5MR status classification (on-track/off-track).\n"
                    "  - World population prospects including birth projections.\n"
                    "  - Global indicator estimates from the global data flow.\n"
                    "  These files are expected to follow specific structural conventions, which are implicitly assumed during parsing."
    )

    doc.add_paragraph("- **Data Cleaning:**\n"
                    "  - U5MR status labels were standardized (e.g., 'achieved', 'on track') and grouped under unified labels "
                    "('On-track' or 'Off-track').\n"
                    "  - The global data flow required reformatting due to inconsistent headers and missing data entries. "
                    "Rows with incomplete 'Geographic area' or 'Indicator' values were removed.\n"
                    "  - Yearly data columns (e.g., 2022, 2021, etc.) were converted to numeric, treating '-' as missing (NaN)."
    )

    doc.add_paragraph("- **Assumptions Made:**\n"
                    "  - The most recent available estimate across years (2022–2018) is a valid proxy for current coverage.\n"
                    "  - Population weight was based solely on the number of projected births, assuming it's a valid proxy for health service need.\n"
                    "  - Country matching across datasets was assumed to be reliable using ISO3 codes and country names (e.g., 'ISO3Code' vs. 'OfficialName')."
    )

    doc.add_paragraph("- **Merging Datasets:**\n"
                    "  - Datasets were merged on country codes and names. Any mismatch or missing alignment (e.g., differing spellings) may lead to data loss."
    )

    doc.add_paragraph("- **Analysis and Visualization:**\n"
                    "  - A population-weighted average coverage was computed for each combination of U5MR status and indicator.\n"
                    "  - Visualization was done using a seaborn barplot with color-coded indicators and status categories."
    )


    # Add plot image
    doc.add_picture(image_path, width=Inches(6))

    doc.add_heading("Results Table", level=1)
    if table_data is not None:
        table = doc.add_table(rows=1, cols=len(table_data.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(table_data.columns):
            hdr_cells[i].text = str(column)

        for _, row in table_data.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    else:
        doc.add_paragraph("No data available to display in the results table.")
    doc.add_paragraph(
        "The data indicates a clear improvement in maternal healthcare coverage from the off-track to on-track status. "
        "Specifically, the percentage of women receiving at least four antenatal care visits increased from 56.52% to 75.92%, "
        "while the proportion of deliveries attended by skilled health personnel rose from 69.38% to 92.72%. "
        "These improvements suggest significant progress in access to and utilization of essential maternal health services, "
        "which are critical for reducing maternal and newborn mortality rates."
    )
    doc.add_paragraph(
        "Position I applied for: Household Survey Data Analyst Consultant - Req. #581656"
    )
    doc.save(output_path)
    print(f"✅ DOCX report saved to {output_path}")


generate_docx_report(image_path="src/consultancy_assessment/documentation/population_weighted_coverage.png", output_path="src/consultancy_assessment/documentation/consultancy_assessment_report.docx")